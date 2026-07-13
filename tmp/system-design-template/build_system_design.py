from copy import deepcopy
from pathlib import Path
from shutil import copy2

from docx import Document

ROOT = Path(r"C:\Users\isaha\Yeni klasör (9)\otopolik-store")
REFERENCE = Path(r"C:\Users\isaha\.codex\plugins\cache\openai-curated-remote\openai-templates\0.1.0\skills\artifact-template-system-design\assets\reference.docx")
OUTPUT = ROOT / "deliverables" / "Otopolik-Store-System-Design.docx"


def replace_paragraph(paragraph, text):
    """Replace text while retaining the source paragraph and run formatting."""
    first_run = paragraph.runs[0] if paragraph.runs else None
    source_rpr = deepcopy(first_run._element.rPr) if first_run and first_run._element.rPr is not None else None
    paragraph.clear()
    run = paragraph.add_run(text)
    if source_rpr is not None:
        run._element.get_or_add_rPr().append(source_rpr)


def replace_cell(cell, text):
    paragraph = cell.paragraphs[0]
    replace_paragraph(paragraph, text)
    for extra in cell.paragraphs[1:]:
        extra.clear()


def set_row(table, row_index, values):
    for cell, value in zip(table.rows[row_index].cells, values):
        replace_cell(cell, value)


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    copy2(REFERENCE, OUTPUT)
    doc = Document(OUTPUT)

    paragraphs = {
        8: "OTO POLİK",
        9: "E-Ticaret Vitrini ve Sipariş Talebi Akışı - Sistem Tasarımı",
        22: (
            "OTO POLİK, araca özel EVA paspaslarını tanıtan bir Next.js App Router vitrini sunar. "
            "Müşteriler ürünleri veya oluşturucuyu kullanarak sepet oluşturur; sipariş talebi ödeme sayfasında alınır, "
            "mümkünse Supabase'e kaydedilir ve WhatsApp üzerinden onay için gönderilir."
        ),
        23: (
            "Tasarım, katalog verisinin uygulama kodunda tutulduğu ve sepetin tarayıcı localStorage'ında yaşadığı mevcut "
            "uygulamayı kapsar. Kartlı ödeme, doğrulanmış işletme iletişim bilgileri ve üretim ortamı izleme henüz kapsam dışıdır."
        ),
        28: (
            "Mevcut satış akışı ürün keşfi, araca göre yapılandırma ve sipariş talebi adımlarını tek bir web uygulamasında "
            "birleştirir. Sipariş onayı WhatsApp'ta ilerlediğinden, mağaza kaydı ile mesajlaşma arasındaki hata görünürlüğü "
            "ve işletme ayarlarının doğruluğu başlıca operasyonel risklerdir."
        ),
        29: (
            "Sistem sınırı tarayıcıda çalışan Next.js arayüzü, localStorage sepeti, isteğe bağlı Supabase sipariş kaydı ve "
            "WhatsApp sipariş bağlantısından oluşur. Temel ilke: müşteri iletişim bilgileri doğrulandıktan sonra sipariş özeti "
            "tutarlı biçimde kullanıcıya ve onay kanalına aktarılmalıdır."
        ),
        33: "Figure 1. OTO POLİK mantıksal mimarisi: Next.js arayüzü -> localStorage sepeti -> Supabase sipariş kaydı + WhatsApp onayı.",
        40: "Müşteri ana sayfa, ürün listesi veya oluşturucudan ürün seçer; ürün, renk/yapılandırma ve miktar bilgisi sepete girer.",
        41: "Ödeme formu ad soyad, telefon, şehir ve adres alanlarını istemci tarafında doğrular. Kimlik doğrulama veya yetkilendirme uygulanmaz.",
        42: "Sepet localStorage'dan yüklenir; ürün fiyatları ve kargo eşiği uygulama içi ürün ve site ayarlarından hesaplanır.",
        43: "Kargo bedeli ile genel toplam hesaplanır; müşteri WhatsApp onayı veya kapıda ödeme tercihini belirtir.",
        44: "Sipariş verisi Supabase yapılandırılmışsa orders tablosuna yazılmaya çalışılır. Yazma hatası konsola kaydedilir ve kullanıcı akışı sürer.",
        45: "Başarılı veya başarısız kayıt denemesinden sonra WhatsApp sipariş bağlantısı yeni pencerede açılır; harici mesajlaşma sonucuna geri çağrı yoktur.",
        46: "Tarayıcı sepeti temizlenir ve teşekkür sayfasına gidilir. Uygulamada tanımlı metrik, iz veya uyarı altyapısı bulunmamaktadır.",
        54: "Gerekli alanlar boş bırakılamaz; telefon biçimi ve adres uzunluğu istemci tarafında kontrol edilir.",
        55: "Siparişler şu an istemci tarafından benzersiz kimlik üretilmeden yazılır; tekrar gönderim için sunucu tarafı idempotency kuralı yoktur.",
        56: "Sipariş satırları, fiyatlar, kargo, müşteri iletişim bilgileri ve ödeme tercihi kayıtta taşınır.",
        57: "Supabase kaydı yapılandırılmışsa operasyonel sipariş kaydıdır; WhatsApp görüşmesi sipariş onayının fiili kanalıdır.",
        58: "Uygulama içi veri sözleşmesi src/app/odeme/page.tsx ve src/lib/types.ts dosyalarında tanımlıdır.",
        63: (
            "Sepet aynı slug ve renk bileşimini tek satırda birleştirir. Sipariş gönderiminde ağ veya Supabase hatası kullanıcıyı "
            "engellemez; bu nedenle aynı talebin tekrarlanması mümkündür. WhatsApp bağlantısı açıldıktan sonra teslim edilmiş bir "
            "mesaj garantisi veya geri oynatma mekanizması yoktur."
        ),
        67: "Uygulama herkese açık vitrindir; müşteri hesabı, oturum açma ve yönetim işlemleri için yetkilendirme bu tasarımda belirtilmemiştir.",
        68: "Ad, telefon, şehir ve adres sipariş talebinde toplanır. Bu alanlar hata kayıtlarına veya analitiğe gereksiz biçimde yazılmamalıdır.",
        69: "Supabase URL ve anonim anahtar NEXT_PUBLIC ortam değişkenleriyle sağlanır. Gizli servis anahtarları istemci paketine konulmamalıdır.",
        70: "Yönetim, tekrar gönderim ve hata ayıklama işlemleri yalnızca doğrulanmış işletme araçlarında yapılmalıdır; mevcut uygulama için süreç tanımlanmalıdır.",
        71: "Saklama, silme ve gizlilik metinleri yayına almadan önce işletmeye özel olarak tamamlanmalıdır.",
        81: "Supabase orders tablosunda benzersiz sipariş kimliği ve tekrar gönderim koruması nasıl uygulanacak?",
        82: "WhatsApp onayından sonra sipariş durumunu güncelleyecek işletme süreci veya yönetim aracı ne olacak?",
        83: "Kartlı ödeme sağlayıcısı eklendiğinde ödeme, sipariş kaydı ve WhatsApp onayı hangi sırayla yürütülecek?",
        84: "Gerçek işletme iletişim bilgileri, fiyatlar, uyumluluk verisi ve hukuki metinler ne zaman doğrulanacak?",
        87: (
            "Mevcut mimari, WhatsApp onaylı sipariş talebi için korunmalıdır. İlk kilometre taşı gerçek işletme ayarları, ürün verisi "
            "ve hukuki metinlerin doğrulanmasıdır. Ardından sipariş kaydı, hata görünürlüğü ve idempotency kuralları test edilerek "
            "sınırlı üretim kullanımına alınmalıdır; kartlı ödeme daha sonra ayrı bir entegrasyon olarak eklenmelidir."
        ),
    }
    for index, text in paragraphs.items():
        replace_paragraph(doc.paragraphs[index], text)

    tables = doc.tables
    set_row(tables[0], 0, ["STATUS", "Draft", "", "OWNER", "OTO POLİK", "", "LAST UPDATED", "11 July 2026"])
    set_row(tables[1], 0, ["Authors", "Project repository"])
    set_row(tables[1], 1, ["Reviewers", "Not specified"])
    set_row(tables[1], 2, ["Related docs", "README.md; src/app/odeme/page.tsx"])
    set_row(tables[1], 3, ["Scope", "Customer web journey from product discovery through WhatsApp order confirmation."])

    goals = [
        ["Present vehicle-specific EVA products and a configurator.", "Implement card payments in this phase."],
        ["Keep cart contents available across browser refreshes.", "Provide authenticated customer accounts."],
        ["Validate essential order-contact data before handoff.", "Guarantee WhatsApp delivery or confirmation."],
        ["Allow a graceful checkout path when Supabase is unavailable.", "Provide production SLOs or alerting without an observability service."],
    ]
    for i, row in enumerate(goals, 1):
        set_row(tables[2], i, row)

    components = [
        ["Next.js App Router", "Renders pages, product discovery, configurator and checkout UI.", "Application source and public media", "Static build failure prevents deployment."],
        ["Cart store", "Keeps cart lines and quantities in the browser.", "localStorage: otopolik-cart", "Starts empty when storage is unavailable or malformed."],
        ["Checkout page", "Validates contact data, totals the order and starts handoff.", "React state; cart snapshot", "Stops on validation errors; proceeds if optional order storage fails."],
        ["Supabase client", "Attempts to persist orders when public configuration is present.", "Supabase orders table", "Logs the failure and continues to WhatsApp handoff."],
        ["WhatsApp link", "Opens a prefilled order-confirmation message.", "wa.me URL", "No delivery acknowledgement is available to the app."],
    ]
    for i, row in enumerate(components, 1):
        set_row(tables[3], i, row)

    fields = [
        ["full_name", "string", "Yes", "Customer full name; minimum 3 characters."],
        ["phone", "string", "Yes", "Turkish mobile number format validated in the browser."],
        ["city", "string", "Yes", "Delivery city entered by the customer."],
        ["address", "string", "Yes", "Delivery address; minimum 10 characters."],
        ["payment_method", "enum", "Yes", "whatsapp or kapida in the current checkout."],
        ["items", "array", "Yes", "Cart lines with slug, name, price, quantity, color and image."],
        ["order_total", "number", "Yes", "Subtotal plus calculated shipping cost."],
    ]
    for i, row in enumerate(fields, 1):
        set_row(tables[4], i, row)

    consistency = [
        ["Duplicate form submission", "Each submission can create another handoff.", "No idempotency key or server-side deduplication exists."],
        ["Supabase write fails", "Checkout continues and opens WhatsApp.", "The current customer handoff is independent of persistence."],
        ["WhatsApp cannot open", "The browser may block or fail to open the external link.", "No callback or retry channel is implemented."],
        ["Configuration changes", "New client-side pricing and site settings apply on the next load.", "No order configuration version is stored."],
    ]
    for i, row in enumerate(consistency, 1):
        set_row(tables[5], i, row)

    operations = [
        ["Checkout completion", "Not yet defined", "Not specified", "Recommended before launch"],
        ["Order persistence failure", "Not yet defined", "Not specified", "Recommended before launch"],
        ["WhatsApp handoff rate", "Not observable in the current application", "Not specified", "Recommended before launch"],
        ["Configuration drift", "Verify siteConfig and product data before deploy", "Project owner", "Required"],
        ["Order integrity", "Sample submitted orders against WhatsApp conversations", "Project owner", "Required"],
        ["Rollout constraint: verify real business contact details, prices, product compatibility and legal texts before public release.", "", "", ""],
    ]
    for i, row in enumerate(operations, 1):
        set_row(tables[6], i, row)

    alternatives = [
        ["WhatsApp-only orders", "Small implementation surface.", "Loses structured order storage when manual handling is needed."],
        ["Supabase-required checkout", "Makes every order persistence attempt mandatory.", "Would block the current fallback to WhatsApp during service failure."],
        ["Immediate card payments", "Provides direct payment capture.", "Marked as a future integration in the repository."],
        ["Server-side cart", "Could support cross-device sessions.", "Customer accounts and server cart APIs are outside current scope."],
    ]
    for i, row in enumerate(alternatives, 1):
        set_row(tables[7], i, row)

    milestones = [
        ["M1", "Verify business and catalog configuration", "Real contacts, prices, compatibility and legal text reviewed."],
        ["M2", "Strengthen order persistence", "Order identity and duplicate-submission policy documented and tested."],
        ["M3", "Limited production release", "Checkout, WhatsApp handoff and order review process exercised."],
        ["M4", "Payment integration decision", "Provider, failure flow and reconciliation design approved."],
    ]
    for i, row in enumerate(milestones, 1):
        set_row(tables[8], i, row)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
